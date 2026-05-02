"""
doc_generator.py
VOC 분석 결과를 Word(.docx) 문서로 생성합니다.
"""

import time
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 색상 ─────────────────────────────────────────────────────
BLUE   = RGBColor(0x14, 0x28, 0xA0)
DARK   = RGBColor(0x1A, 0x1A, 0x2E)
GRAY   = RGBColor(0x44, 0x44, 0x55)
LGRAY  = RGBColor(0x99, 0x99, 0xAA)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def _shading(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def _hdr_cell(text: str):
    from docx.table import _Cell
    # 반환은 caller에서 처리
    return text


def _add_header_row(table, headers: List[str]) -> None:
    row = table.rows[0]
    for i, h in enumerate(headers):
        c = row.cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        _shading(c, "1428A0")


def _add_body_row(table, *cols: str, center_from: int = 1) -> None:
    row = table.add_row()
    for i, text in enumerate(cols):
        c = row.cells[i]
        c.text = str(text or "")
        if i >= center_from:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in c.paragraphs[0].runs:
            run.font.size = Pt(10)


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = BLUE
    p.paragraph_format.border_bottom = True


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DARK


def _body(doc: Document, text: str, indent: bool = False) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(f"• {text}")
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY


def generate_docx(
    voc_list:    list,
    analysis:    dict,
    srs_text:    str,
    product_name: str = "삼성 갤럭시",
    version:     str = "1.0",
    author:      str = "제품기획팀",
    output_dir:  str = "./output",
) -> str:
    """Word 문서 생성 후 파일 경로 반환"""
    from collections import Counter

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    date = time.strftime("%Y년 %m월 %d일")

    # ── 카테고리/소스/감성 통계 ──────────────────────────────
    def get_attr(v, attr):
        return getattr(v, attr) if hasattr(v, attr) else v.get(attr, "")

    cats = Counter(get_attr(v, "category") for v in voc_list)
    srcs = Counter(get_attr(v, "source")   for v in voc_list)
    snts = Counter(get_attr(v, "sentiment") for v in voc_list)

    # ── 표지 ─────────────────────────────────────────────────
    doc.add_paragraph("\n")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(product_name)
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = BLUE

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("VOC 기반 소프트웨어 요구사항명세서")
    r2.font.size = Pt(17); r2.font.color.rgb = DARK

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run(f"v{version}  |  {date}  |  {author}  |  Gemma 3n E2B-it")
    r3.font.size = Pt(11); r3.font.color.rgb = LGRAY

    doc.add_page_break()

    # ── 1. 수집 현황 ─────────────────────────────────────────
    _h1(doc, "1. VOC 수집 현황")

    t_sum = doc.add_table(rows=1, cols=4)
    t_sum.style = "Table Grid"
    _add_header_row(t_sum, ["총 수집", "부정 의견", "긍정 의견", "수집 채널"])
    _add_body_row(t_sum,
        f"{len(voc_list):,}건",
        f"{snts['negative']}건 ({round(snts['negative']/max(len(voc_list),1)*100)}%)",
        f"{snts['positive']}건 ({round(snts['positive']/max(len(voc_list),1)*100)}%)",
        f"{len(srcs)}개",
        center_from=0,
    )
    doc.add_paragraph()

    _h2(doc, "카테고리별 분포")
    t_cat = doc.add_table(rows=1, cols=3)
    t_cat.style = "Table Grid"
    _add_header_row(t_cat, ["카테고리", "건수", "비율"])
    for cat, cnt in cats.most_common():
        _add_body_row(t_cat, cat, f"{cnt}건", f"{round(cnt/len(voc_list)*100)}%")
    doc.add_paragraph()

    _h2(doc, "채널별 현황")
    t_src = doc.add_table(rows=1, cols=3)
    t_src.style = "Table Grid"
    _add_header_row(t_src, ["수집 채널", "건수", "비율"])
    for src, cnt in srcs.most_common():
        _add_body_row(t_src, src, f"{cnt}건", f"{round(cnt/len(voc_list)*100)}%")
    doc.add_paragraph()

    # ── 2. 종합 분석 ─────────────────────────────────────────
    if analysis.get("executive_summary"):
        _h1(doc, "2. 종합 분석 요약")
        _body(doc, analysis["executive_summary"])

        if analysis.get("key_insights"):
            _h2(doc, "핵심 인사이트")
            for ins in analysis["key_insights"]:
                _bullet(doc, ins)
        doc.add_paragraph()

    # ── 3. 핵심 이슈 ─────────────────────────────────────────
    if analysis.get("critical_issues"):
        _h1(doc, "3. 핵심 이슈 분석")
        t_iss = doc.add_table(rows=1, cols=4)
        t_iss.style = "Table Grid"
        _add_header_row(t_iss, ["이슈", "카테고리", "빈도", "영향도"])
        for iss in analysis["critical_issues"]:
            _add_body_row(t_iss,
                iss.get("title", ""), iss.get("category", ""),
                iss.get("frequency", ""), iss.get("impact", ""),
            )
        doc.add_paragraph()

        # 이슈 상세
        _h2(doc, "이슈 상세 설명")
        for iss in analysis["critical_issues"]:
            p = doc.add_paragraph()
            r = p.add_run(f"▶ {iss.get('title', '')}")
            r.font.bold = True; r.font.color.rgb = BLUE; r.font.size = Pt(11)
            _body(doc, iss.get("description", ""), indent=True)

    # ── 4. 기능 요구사항 ──────────────────────────────────────
    if analysis.get("requirements"):
        doc.add_page_break()
        _h1(doc, "4. 기능 요구사항")

        groups: dict = {}
        for req in analysis["requirements"]:
            groups.setdefault(req.get("category", "기타"), []).append(req)

        for idx, (cat, reqs) in enumerate(groups.items(), 1):
            _h2(doc, f"4.{idx} {cat}")
            for req in reqs:
                p = doc.add_paragraph()
                r = p.add_run(f"[{req.get('id','?')}] {req.get('title','?')}")
                r.font.bold = True; r.font.size = Pt(11)
                pri = req.get("priority", "?")
                r2 = p.add_run(f"  [{pri}]")
                r2.font.bold = True; r2.font.size = Pt(10)
                r2.font.color.rgb = RGBColor(0xCC,0x00,0x00) if pri=="필수" else (
                    RGBColor(0xE6,0x51,0x00) if pri=="권장" else RGBColor(0x2E,0x7D,0x32))

                _body(doc, f"설명: {req.get('description','')}", indent=True)
                _body(doc, f"사용자 스토리: {req.get('user_story','')}", indent=True)
                if req.get("acceptance_criteria"):
                    _body(doc, "검증 기준:", indent=True)
                    for c in req["acceptance_criteria"]:
                        _bullet(doc, c)

    # ── 5. AI 생성 SRS 전문 ───────────────────────────────────
    if srs_text:
        doc.add_page_break()
        _h1(doc, "5. 요구사항명세서 전문 (Gemma 3n E2B-it 생성)")
        for line in srs_text.split("\n"):
            t = line.strip()
            if not t:
                doc.add_paragraph()
                continue
            if t.startswith("# "):   _h1(doc, t[2:])
            elif t.startswith("## "): _h2(doc, t[3:])
            elif t.startswith("### "): _h2(doc, t[4:])
            elif t.startswith(("- ", "• ")): _bullet(doc, t[2:])
            else: _body(doc, t)

    # ── 6. 로드맵 ─────────────────────────────────────────────
    if analysis.get("roadmap"):
        doc.add_page_break()
        _h1(doc, "6. 개선 로드맵")
        for phase in analysis["roadmap"]:
            _h2(doc, phase.get("phase", ""))
            for item in phase.get("items", []):
                _bullet(doc, item)

    # ── 7. 승인 ──────────────────────────────────────────────
    doc.add_paragraph()
    _h1(doc, "7. 검토 및 승인")
    t_ap = doc.add_table(rows=1, cols=4)
    t_ap.style = "Table Grid"
    _add_header_row(t_ap, ["구분", "성명", "서명", "일자"])
    _add_body_row(t_ap, "작성자", "", "", time.strftime("%Y-%m-%d"))
    _add_body_row(t_ap, "검토자", "", "", "")
    _add_body_row(t_ap, "승인자", "", "", "")

    # ── 저장 ─────────────────────────────────────────────────
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    ts       = time.strftime("%Y%m%d_%H%M%S")
    filename = f"galaxy_voc_srs_{ts}.docx"
    filepath = Path(output_dir) / filename
    doc.save(str(filepath))
    return str(filepath)
